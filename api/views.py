# views.py (full file)

from dotenv import load_dotenv
import os
import json
import base64
import io
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import OxmlElement as SharedOxmlElement
import google.generativeai as genai
from google.api_core.exceptions import GoogleAPIError
import time
import logging
import traceback
import textract
import pdfplumber
from PyPDF2 import PdfReader

load_dotenv()  # Load environment variables from .env file
API_KEY = os.environ.get("GOOGLE_API_KEY", "")
if not API_KEY:
    raise ValueError("GOOGLE_API_KEY is not set. Please configure it in the .env file or environment.")
genai.configure(api_key=API_KEY)
MODEL_NAME = "gemini-2.0-flash"  # Use available model
MAX_RETRIES = 3

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# JSON Schema remains the same
WB_CV_SCHEMA = {
    "type": "object",
    "properties": {
        "name": {
            "type": "string",
            "description": "Full name of the expert"
        },
        "expert_contact_information": {
            "type": "object",
            "properties": {
                "phone": {"type": "string",
                          "description": "Phone or mobile number of the expert"},
                "email": {"type": "string",
                          "description": "Email address of the expert"}
            },
            "required": ["phone", "email"]
        },
        "proposed_position": {
            "type": "string",
            "description": "Position title and number"
        },
        "employer": {
            "type": "string",
            "description": "Current employer"
        },
        "date_of_birth": {
            "type": "string",
            "description": "Date of birth"
        },
        "nationality": {
            "type": "string",
            "description": "Country of citizenship/residence"
        },
        "education": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "school_university": {"type": "string"},
                    "degree": {"type": "string"},
                    "date_obtained": {"type": "string"}
                },
                "required": ["school_university", "degree", "date_obtained"]
            }
        },
        "membership_in_professional_associations": {
            "type": "string",  # Changed to string to match frontend
            "description": "Memberships in professional associations as a multi-line string (e.g., join with '\n'). Empty if not found."
        },
        "publications": {
            "type": "string",  # Changed to string to match frontend
            "description": "Publications as a multi-line string (e.g., join with '\n'). Empty if not found."
        },
        "other_training": {
            "type": "string",
            "description": "All relevant training information copied exactly as in CV and must be given result by listing of all trainings. Empty string if not present."
        },
        "countries_experience": {  # Added missing property
            "type": "string",
            "description": "List of countries with experience, comma-separated. Empty if not found."
        },
        "languages": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "language": {"type": "string"},
                    "speaking": {"type": "string"},
                    "reading": {"type": "string"},
                    "writing": {"type": "string"}
                },
                "required": ["language", "speaking", "reading", "writing"]
            }
        },
        "employment_record": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "from": {"type": "string"},
                    "to": {"type": "string"},
                    "employer": {"type": "string", "description": "Name of the employer (mandatory)"},
                    "position": {"type": "string", "description": "Title of the position held just title of the position nothing else exactly in cv (mandatory)"},
                    "location": {"type": "string", "description": "City and country of the employer"},
                    "summary_of_activities": {"type": "string", "description": "Brief description of the main activities and responsibilities"},
                    "for_references": {"type": "string", "description": "Whether the employer can be contacted for references just find the reference and no give answer"},
                    "name": {"type": "string", "description": "Name of the person which is the reference in cv for spacific company or project (mandatory)"},
                    "designation": {"type": "string", "description": "Designation of the expert in this project find it important data or position held (mandatory)"},
                    "telephone": {"type": "string", "description": "find the reference contact Phone number, mobile number, contact number (mandatory)"},
                    "email": {"type": "string", "description": "Email address of the expert or mail address (mandatory)"}
                },
                "required": ["from", "to", "employer", "position", "for_references", 
                             "name", "designation", "telephone", "email", "location", 
                             "summary_of_activities"]
            }
        },
        "detailed_tasks": {
            "type": "array",
            "items": {"type": "string"},
            "description": "no need any result send empty "
        },
        "work_undertaken": {
            "type": "array",
            "items": {
                "type": "object",
                "description": "Find project by Project name/date/Details of each assignment/project undertaken, Mainly find Name of project and put all projects one by one with all project details/date/main features/position held/activities undertaken exactly as in cv. Check cv very well for all projects no miss any project.",
                "properties": {
                    "name": {"type": "string", "description": "Name of the assignment/project find from the CV and copy exactly as in cv. must be given full project name by any change not found in cv then give empty string"},
                    "year": {"type": "string", "description": "Date of the assignment/project mention start date to end date or Present/till date. Given result must be exactly as in cv"},
                    "location": {"type": "string"},
                    "client": {"type": "string"},
                    "main_features": {"type": "string", "description": "Main features of the project, try to find this from the cv and copy exactly as in cv every project must have main features."},
                    "position_held": {"type": "string", "description": "Position held during the project. this must be given exactly as in cv."},
                    "activities": {"type": "string", "description": "Description of the main activities undertaken, responsibilities carried out during the project. Try to find this from the CV and copy exactly as in CV."}
                },
                "required": ["name", "year", "location", "client", "main_features", 
                             "position_held", "activities"]  # Removed "summary"
            }
        },
        "worked_for_world_bank": {
            "type": "string",
            "description": "Details of World Bank work experience or 'No'"
        }
    },
    "required": [  # Uncommented and corrected
        "name", "expert_contact_information", "proposed_position", "employer", "date_of_birth", "nationality",
        "education", "membership_in_professional_associations", "publications", "other_training", "countries_experience",
        "languages", "employment_record", "detailed_tasks", "work_undertaken",
        "worked_for_world_bank"
    ]
}

def extract_text(base64_content, filename):
    try:
        ext = filename.lower().split('.')[-1]
        binary = base64.b64decode(base64_content)

        full_text = []

        if ext == 'docx':
            doc = Document(io.BytesIO(binary))
            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(' | '.join(row_text))
            # Extract bullet points
            for para in doc.paragraphs:
                if para.style.name.startswith('List') or para.text.strip().startswith(('*', '-', '•')):
                    full_text.append(para.text.strip())
        
        elif ext == 'pdf':
            with pdfplumber.open(io.BytesIO(binary)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        lines = text.split('\n')
                        for line in lines:
                            stripped = line.strip()
                            if stripped:
                                full_text.append(stripped)
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                            if row_text:
                                full_text.append(' | '.join(row_text))
        
        elif ext == 'txt':
            full_text.append(binary.decode('utf-8').strip())
        
        elif ext == 'doc':
            text = textract.process(io.BytesIO(binary))
            full_text.append(text.decode('utf-8').strip())
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        extracted_text = "\n".join(full_text)
        logger.info(f"Extracted text from {ext.upper()} (length: {len(extracted_text)}): {extracted_text[:500]}...")  # Log snippet
        return extracted_text
    except Exception as e:
        logger.error(f"Error extracting text from file: {e}")
        return ""

@csrf_exempt
@require_http_methods(["POST"])
def process_cv_view(request):
    try:
        data = json.loads(request.body)
        file_content_base64 = data.get('file_content')
        filename = data.get('filename')

        if not file_content_base64 or not filename:
            return JsonResponse({'success': False, 'message': 'File content or filename is missing.'}, status=400)

        cv_text = extract_text(file_content_base64, filename)
        if not cv_text:
            return JsonResponse({'success': False, 'message': 'Failed to extract text from file.'}, status=400)

        model = genai.GenerativeModel(MODEL_NAME)
        prompt = (
            """You are an expert in extracting data from CVs strictly according to the World Bank FORM TECH-6 template. 

                Your sole task is to parse the provided CV text and extract ONLY the information that is explicitly stated. Do NOT invent, assume, infer, modify, or add any data. If a value or section is not present, missing, or unclear, you MUST return an empty string ('') for strings, and an empty array ([]) for lists or arrays in the output.

                Every field listed in the JSON schema below must be present in your output, even if the corresponding data is missing from the CV. If data is not found for a particular field, provide the exact required empty value for that field ('' or []). Under NO circumstances should you omit any fields or return a partial/malformed object.

                Preserve original wording, phrasing, capitalization, and formatting as much as possible. Normalize dates only to 'YYYY', 'YYYY-MM', or 'YYYY-MM-DD' formats if they are clearly dates; otherwise, leave as-is.

                The CV text may contain:
                - Plain paragraphs: e.g., 'Name: John Doe'
                - Table rows separated by ' | ', e.g., 'University X | MSc | 2020'
                - Bullet points starting with '*', '-', or '•'

                Identify sections by number (1., 2., etc.) or headings (Education, Languages, etc.). Extract all required information as arrays in the order they appear, sorting employment_record in reverse chronological order if dates can be parsed.

                Fields to extract (all must be present in output, always fill with either actual data or empty as instructed above):
                1. name: Extract full name after 'Name' or similar; '' if not found.
                2. expert_contact_information: Object with phone (extract phone/mobile, '' if missing) and email (extract email, '' if missing).
                3. proposed_position: Extract after 'Proposed Position' or similar; '' if not found.
                4. employer: Extract current/primary employer; '' if not found.
                5. date_of_birth: Extract birth date exactly; '' if not found.
                6. nationality: Extract nationality; '' if not found.
                7. education: Array of objects from education section/table, or [] if not found. For each: school_university, degree, date_obtained (all fields '' if missing).
                8. membership_in_professional_associations: Multi-line string; '' if not found.
                9. publications: Multi-line string; '' if not found.
                10. other_training: Training text, combine lines as needed; '' if not found.
                11. countries_experience: Comma-separated string; '' if not found.
                12. languages: Array of objects from languages section, or [] if not found. For each: language, speaking, reading, writing (all fields '' if missing).
                13. employment_record: Array of objects, or [] if not found. See schema for required subfields; use '' for any missing string field.
                14. detailed_tasks: Array of strings; [] if not found.
                15. work_undertaken: Array of project objects, or [] if not found. All subfields as above.
                16. worked_for_world_bank: Extract answer to World Bank work question; default to 'No' if absent.

                Below is the required JSON schema. Output ONLY a JSON object with EVERY field as described; do not output any additional explanation, text, or formatting."""
                "CV text:\n"
            f"{cv_text}"
        )

        config = genai.types.GenerationConfig(
            response_mime_type="application/json",
            response_schema=WB_CV_SCHEMA
        )

        last_error = None
        for attempt in range(MAX_RETRIES):
            try:
                response = model.generate_content([prompt], generation_config=config)
                json_text = response.text
                parsed_cv_data = json.loads(json_text)
                logger.info(f"Successfully processed CV data: {parsed_cv_data}")
                return JsonResponse({'success': True, 'cv_data': parsed_cv_data})
            except GoogleAPIError as e:
                last_error = e
                logger.error(f"Gemini API Error on attempt {attempt + 1}: {e}")
                if attempt < MAX_RETRIES - 1:
                    delay = 2 ** attempt
                    logger.info(f"Retrying in {delay} seconds...")
                    time.sleep(delay)

        error_message = f"Failed to communicate with AI service after {MAX_RETRIES} retries. Last error: {last_error}"
        logger.error(error_message)
        return JsonResponse({"error": error_message}, status=503)

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON body.'}, status=400)
    except Exception as e:
        logger.error(f"Server error during processing: {e}")
        return JsonResponse({'success': False, 'message': f'Server error during processing: {e}'}, status=500)

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_paragraph_background(paragraph, fill_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    pPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:' + edge
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)


@csrf_exempt
@require_http_methods(["POST"])
def generate_docx_view(request):
    try:
        data = json.loads(request.body)
        cv_data = data.get('cv_data')

        if not cv_data:
            return JsonResponse({'success': False, 'message': 'CV data is missing.'}, status=400)

        docx_file_stream = (cv_data)
        
        response = HttpResponse(
            docx_file_stream.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename=Formatted_WB_CV_{cv_data.get("name", "Document").replace(" ", "_")}.docx'
        
        return response

    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'message': 'Invalid JSON body.'}, status=400)
    except Exception as e:
        logger.error(f"DOCX Generation Server Error: {str(e)} with traceback: {traceback.format_exc()}")
        return JsonResponse({'success': False, 'message': f'Server error during DOCX generation: {str(e)}'}, status=500)